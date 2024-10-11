# Step 1: Import Libraries and Define Helper Functions
import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.io as pio
from io import BytesIO
from docx import Document
from docx.shared import Inches
import base64

# Set Plotly theme to 'plotly_white'
pio.templates.default = "plotly_white"

# Helper function to create download links
def download_link(object_to_download, download_filename, download_link_text):
    b64 = base64.b64encode(object_to_download).decode()
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{download_filename}">{download_link_text}</a>'

# Function to save Plotly figures as images
def save_plotly_fig_as_image(fig, filename):
    fig.write_image(filename)

# Function to list all sheets in the uploaded Excel file
def list_excel_sheets(uploaded_file):
    if uploaded_file is not None:
        xls = pd.ExcelFile(uploaded_file)
        return xls.sheet_names
    return []
# Step 2: Handle File Uploads for Course Details
def handle_file_uploads():
    st.header("Upload Course Details File")
    uploaded_course_file = st.file_uploader("Upload Course Details Excel File", type="xlsx", key="course_file")

    course_details = None
    if uploaded_course_file:
        course_sheets = list_excel_sheets(uploaded_course_file)
        selected_sheet = st.selectbox("Select the sheet to load:", course_sheets, key="course_sheet_select")
        course_details = pd.read_excel(uploaded_course_file, sheet_name=selected_sheet)

        # Verify required columns in Course Details
        required_columns = ['Course Name', 'Country', 'University Name', 'Tuition Fees (INR)', 'Living Expenses (INR)']
        for col in required_columns:
            if col not in course_details.columns:
                st.error(f"Required column '{col}' is missing in Course Details. Please ensure that your file contains this column.")
                return None, []

        st.success("Course Details uploaded successfully!")

    return course_details
# Step 3: Handle Collection Data Uploads
def handle_collection_uploads():
    st.header("Upload Collection Data File")
    uploaded_collection_file = st.file_uploader("Upload Collection Excel File", type="xlsx", key="collection_file")

    collection_data_list = []
    
    if uploaded_collection_file:
        collection_sheets = list_excel_sheets(uploaded_collection_file)
        
        for sheet_name in collection_sheets:
            collection_data = pd.read_excel(uploaded_collection_file, sheet_name=sheet_name)
            required_columns = ['Course Name', 'Country', 'Collection Name', 'University Name', 'Tuition Fees', 'Acceptance Rate', 'Application Link', 'Ranking', 'Agency Name', 'Stream']
            missing_columns = [col for col in required_columns if col not in collection_data.columns]
            
            if missing_columns:
                st.warning(f"The following columns are missing in the '{sheet_name}' sheet: {', '.join(missing_columns)}. Please correct the values in the file.")
                continue
            
            collection_data_list.append({
                "sheet_name": sheet_name,
                "data": collection_data
            })

    return collection_data_list
# Step 4: Read Third Excel File
def read_third_excel_file(uploaded_file):
    if uploaded_file:
        xls = pd.ExcelFile(uploaded_file)

        # Verify the required sheets exist
        expected_sheets = ['Tuition Fees', 'Living Expenses', 'Important Deadlines']
        missing_sheets = [sheet for sheet in expected_sheets if sheet not in xls.sheet_names]
        
        if missing_sheets:
            st.error(f"The following sheets are missing in the uploaded file: {', '.join(missing_sheets)}.")
            return None, None, None

        # Read each sheet into a DataFrame
        tuition_fees_data = pd.read_excel(xls, sheet_name='Tuition Fees')
        living_expenses_data = pd.read_excel(xls, sheet_name='Living Expenses')
        deadlines_data = pd.read_excel(xls, sheet_name='Important Deadlines')

        # Verify required columns in Tuition Fees
        if 'Tuition Fees (INR)' not in tuition_fees_data.columns:
            st.error("Column 'Tuition Fees (INR)' is missing in Tuition Fees data.")
            return None, None, None

        # Verify required columns in Living Expenses
        if 'Monthly Cost (INR)' not in living_expenses_data.columns and 'Annual Cost (INR)' not in living_expenses_data.columns:
            st.error("Columns 'Monthly Cost (INR)' and 'Annual Cost (INR)' are missing in Living Expenses data.")
            return None, None, None

        return tuition_fees_data, living_expenses_data, deadlines_data
# Step 5: Generate Graph Data
def generate_graph_data(course_details):
    required_columns = ['University Name', 'Tuition Fees (INR)', 'Acceptance Rate (%)']
    for col in required_columns:
        if col not in course_details.columns:
            st.error(f"Required column '{col}' is missing in Course Details. Please ensure that your file contains this column.")
            return pd.DataFrame()  # Return an empty DataFrame in case of error
    
    # Prepare data suitable for visualizations
    graph_data = course_details[required_columns].copy()
    return graph_data
# Step 6: Generate Final Content
def generate_final_content(course_details, graph_data, collection_data_list, course_name, country, tuition_fees_data, living_expenses_data, deadlines_data):
    # Filter course details based on the selected course and country
    selected_course = course_details[(course_details['Course Name'] == course_name) & 
                                      (course_details['Country'] == country)]
    
    if selected_course.empty:
        st.error("No course details found for the selected course and country.")
        return ""

    content = f"# {course_name} in {country}\n\n"
    
    # Overview Section
    content += f"## Overview\n"
    content += (f"The {course_name} in {country} is a prestigious degree attracting international students. "
                f"Top specializations include {selected_course.iloc[0]['Top Specializations']}, and universities like "
                f"{selected_course.iloc[0]['University Name']} are renowned for their research opportunities.\n\n")
    
    # Embed YouTube video after Overview
    youtube_link = selected_course.iloc[0]['Relevant YouTube Video URL']  # Get from the DataFrame
    content += f"You can watch this video for more information: [Watch Video]({youtube_link})\n\n"

    # Cost Section: Ensure tuition fees and living expenses data are valid before accessing
    tuition_fees = 0
    living_expenses = 0
    
    if tuition_fees_data is not None:
        tuition_fees = tuition_fees_data['Tuition Fees (INR)'].sum()  # Calculate total tuition fees
    
    if living_expenses_data is not None:
        living_expenses = living_expenses_data['Monthly Cost (INR)'].sum()  # Calculate total living expenses
    
    content += f"## Cost and Living Expenses\n"
    content += (f"The cost of studying {course_name} in {country} ranges from **{tuition_fees}** for tuition and approximately "
                f"**{living_expenses}** for living expenses per year.\n\n")
    
    # Add a table for Cost and Living Expenses
    content += "| **Type** | **Amount (INR)** |\n"
    content += "|----------|------------------|\n"
    content += f"| Tuition Fees | **{tuition_fees}** |\n"
    content += f"| Living Expenses | **{living_expenses}** |\n\n"

    # Important Deadlines Section
    content += "## Important Deadlines\n"
    if deadlines_data is not None and not deadlines_data.empty:
        content += "| **Event** | **Date** | **Notes** |\n"
        content += "|----------|----------|-----------|\n"
        for _, row in deadlines_data.iterrows():
            content += f"| {row['Event']} | {row['Date']} | {row.get('Notes', 'N/A')} |\n"
    else:
        content += "No important deadlines available.\n\n"

    # Add each collection with H2 headers, brief descriptions, and formatted tables
    unique_collections = {data["sheet_name"]: data for data in collection_data_list}  # Avoid duplicates

    for collection_data in unique_collections.values():
        sheet_name = collection_data["sheet_name"]
        data = collection_data["data"]
        description = collection_data.get("description", "")
        
        # Add Section Header and Description
        content += f"## {sheet_name}\n\n"
        content += f"{description}\n\n"
        
        # Add the collection data as a formatted Markdown table
        content += "| **University Name** | **Tuition Fees (INR)** | **Acceptance Rate (%)** | **Application Link** | **Ranking** |\n"
        content += "|---------------------|------------------------|------------------------|----------------------|-------------|\n"
        
        for _, row in data.iterrows():
            ranking = f"Ranked {row['Ranking']} in {row['Stream']} by {row['Agency Name']}" if pd.notna(row['Ranking']) else 'N/A'
            content += (f"| {row['University Name']} | {row['Tuition Fees']} | {row['Acceptance Rate']}% | "
                        f"[Apply Here]({row['Application Link']}) | {ranking} |\n")
    
    # Job Prospects Section
    content += "\n## Job Prospects and Career Growth\n"
    content += (f"Graduates from {course_name} in {country} often secure high-paying roles with companies like "
                f"{selected_course.iloc[0]['Top Employers']}. Starting salaries average around {selected_course.iloc[0]['Median Salary (USD)']}.\n\n")

    return content

def generate_word_with_images(content, collection_data_list, graph_image_paths):
    doc = Document()
    
    # Add document title
    doc.add_heading(content.split('\n')[0], 0)  
    
    # Add paragraphs from content
    paragraphs = content.split('\n\n')
    for para in paragraphs[1:]:
        doc.add_paragraph(para)
    
    # Add collections (with tables) from the collection_data_list
    for collection_data in collection_data_list:
        sheet_name = collection_data["sheet_name"]
        data = collection_data["data"]
        description = collection_data.get("description", "")
        
        # Add section header and description
        doc.add_heading(f"{sheet_name}", level=2)
        doc.add_paragraph(description)
        
        # Create a table
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'University Name'
        hdr_cells[1].text = 'Tuition Fees'
        hdr_cells[2].text = 'Acceptance Rate'
        hdr_cells[3].text = 'Application Link'
        hdr_cells[4].text = 'Ranking'

        # Populate the table with data
        for _, row in data.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['University Name'])
            row_cells[1].text = str(row['Tuition Fees']) if pd.notna(row['Tuition Fees']) else 'N/A'
            row_cells[2].text = str(row['Acceptance Rate']) + '%' if pd.notna(row['Acceptance Rate']) else 'N/A'
            row_cells[3].text = str(row['Application Link']) if pd.notna(row['Application Link']) else 'N/A'
            row_cells[4].text = f"Ranked {row['Ranking']} in {row['Stream']} by {row['Agency Name']}" if pd.notna(row['Ranking']) else 'N/A'
    
    # Add graphs to the Word document
    doc.add_heading('Graphs and Visualizations', level=1)
    for img_path in graph_image_paths:
        doc.add_picture(img_path, width=Inches(5.0))
    
    # Save the document to memory
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer
# Step 8: Generate HTML Content for Download
def generate_html(content, graph_image_paths):
    html_content = f"<html><head><title>{content.split('\n')[0]}</title></head><body>"
    paragraphs = content.split('\n\n')

    for para in paragraphs:
        if para.startswith("##"):
            html_content += f"<h2>{para[2:].strip()}</h2>"
        elif para.startswith("#"):
            html_content += f"<h1>{para[1:].strip()}</h1>"
        else:
            html_content += f"<p>{para.strip()}</p>"
    
    # Add images in HTML
    for img_path in graph_image_paths:
        html_content += f'<img src="{img_path}" style="width:100%;"/>'  # Embed images in the HTML
    
    html_content += "</body></html>"
    return html_content
# Step 9: Main Application Function to Integrate All Components
def main():
    st.title("University Course Automation with Rankings")
    
    # Step 1: Handle file uploads for course and collection data
    course_details = handle_file_uploads()
    
    # Step 2: Upload Collection Data file
    collection_data_list = handle_collection_uploads()

    # Step 3: Upload Third Excel File
    st.header("Upload University Details File")
    uploaded_details_file = st.file_uploader("Upload University Details Excel File", type="xlsx", key="details_file")

    tuition_fees_data = None
    living_expenses_data = None
    deadlines_data = None
    
    if uploaded_details_file:
        tuition_fees_data, living_expenses_data, deadlines_data = read_third_excel_file(uploaded_details_file)

        # Display Tuition Fees Data
        if tuition_fees_data is not None:
            st.subheader("Tuition Fees Data")
            st.write(tuition_fees_data)

        # Display Living Expenses Data
        if living_expenses_data is not None:
            st.subheader("Living Expenses Data")
            st.write(living_expenses_data)

        # Display Important Deadlines Data
        if deadlines_data is not None:
            st.subheader("Important Deadlines Data")
            st.write(deadlines_data)

    if course_details is not None and collection_data_list:
        # Step 4: Select Course and Country
        course_name = st.selectbox("Select Course Name", course_details['Course Name'].unique())
        country = st.selectbox("Select Country", course_details['Country'].unique())
        
        # Generate graph data for visualization
        graph_data = generate_graph_data(course_details)  # Create graph data based on your requirements
        
        # Generate final content
        generated_content = generate_final_content(
            course_details,
            graph_data,
            collection_data_list,
            course_name,
            country,
            tuition_fees_data,
            living_expenses_data,
            deadlines_data
        )
        
        # Step 5: Display Content
        if generated_content:
            st.subheader("Generated Content")
            st.markdown(generated_content)

            # Step 6: Visualizations
            if graph_data is not None and not graph_data.empty:
                # Create Bar Plot
                bar_fig = px.bar(graph_data, x='University Name', y='Tuition Fees (INR)',
                                title='Tuition Fees by University')
                st.plotly_chart(bar_fig)

                # Create Scatter Plot
                scatter_fig = px.scatter(graph_data, x='University Name', y='Tuition Fees (INR)',
                                        size='Acceptance Rate (%)', hover_name='University Name',
                                        title='Tuition Fees vs Acceptance Rate')
                st.plotly_chart(scatter_fig)

                # Save the plots as images for the Word document
                bar_image_path = "/tmp/bar_plot.png"
                scatter_image_path = "/tmp/scatter_plot.png"
                save_plotly_fig_as_image(bar_fig, bar_image_path)
                save_plotly_fig_as_image(scatter_fig, scatter_image_path)

                # Pass the saved image paths to be included in the Word document later
                graph_image_paths = [bar_image_path, scatter_image_path]
            else:
                graph_image_paths = []

            # Step 7: Download Word and HTML
            word_file = generate_word_with_images(generated_content, collection_data_list, graph_image_paths)
            st.markdown(download_link(word_file.read(), f"{course_name}_complete_with_rankings.docx", "Download Word Document"), unsafe_allow_html=True)
            
            html_content = generate_html(generated_content, graph_image_paths)
            st.markdown(download_link(html_content.encode(), f"{course_name}_complete_content.html", "Download HTML Document"), unsafe_allow_html=True)

if __name__ == "__main__":
    main()
